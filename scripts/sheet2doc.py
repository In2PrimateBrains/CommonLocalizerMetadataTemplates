import pandas as pd
from docx import Document
from docx.shared import Inches

def read_spreadsheet(file_path):
    spreadsheets = pd.read_excel(file_path, None)
    return spreadsheets

def add_bullet_paragraph(document, textBold, textItalics, text, level=0):
    p = document.add_paragraph(style='ListBullet')
    p.paragraph_format.left_indent = Inches(0.25 * level)
    p.add_run(textBold).bold = True
    p.add_run(textItalics).italic = True
    p.add_run(text)

def generate_word_document(df, sheet_name, doc):
    

    doc.add_heading(f"{sheet_name} \n", level=2)

    for _, row in df.iterrows():
        level1, level2, level3 = row["DataLevel1"], row["DataLevel2"], row["DataLevel3"]
        requirement_level = row["RequirementLevel"]
        description, example, allowed_values, comments = row["Description"], row["Example"], row["AllowedValues"], row["Comments"]

        if not pd.isna(level1):
            add_bullet_paragraph(doc, f"{level1} ", f"[{requirement_level}]: ", f"{description}", level=0)

        if not pd.isna(level2):
            add_bullet_paragraph(doc, f"{level2} ", f"[{requirement_level}]: ", f"{description}", level=1)

        if not pd.isna(level3):
            add_bullet_paragraph(doc, f"{level3} ", f"[{requirement_level}]: ", f"{description}", level=2)

        if not pd.isna(example):
            doc.add_paragraph(f"Example: {example}")
        if not pd.isna(allowed_values):
            doc.add_paragraph(f"Allowed Values: {allowed_values}")
        if not pd.isna(comments):
            doc.add_paragraph(f"Note: {comments}")

    


if __name__ == "__main__":
    file_path = "metadata_fields.xlsx"  
    output_file = "metadata_fields.docx"  

    doc = Document()

    sheets = read_spreadsheet(file_path)
    for sheetname in sheets:
        generate_word_document(sheets[sheetname], sheet_name = sheetname, doc=doc)

    doc.save(output_file)
